import os
import ast
from datetime import datetime

# Import analyzers and AI engine
from ai.ai_engine import ask_ai
from utils.file_reader import read_file

# Check if OCR / PaddleOCR is available
try:
    from paddleocr import PaddleOCR
    PADDLE_AVAILABLE = True
except ImportError:
    PADDLE_AVAILABLE = False

def generate_project_report(folder_path, save_directory="reports", fmt="md"):
    """
    Scans a folder, gathers statistics, extracts code metadata (classes, functions, frameworks),
    scans logs for errors, runs OCR on found images, reads documents, and generates a
    master AI project summary. Writes a premium report to the specified directory in the requested format.
    """
    if not os.path.isdir(folder_path):
        print(f"Error: {folder_path} is not a valid directory.")
        return None

    project_name = os.path.basename(os.path.abspath(folder_path))
    if not project_name:
        project_name = "Local Project"

    print(f"\n[Report Generator] Commencing deep analysis of folder: {folder_path}", flush=True)

    # Walk directory and ignore virtualenvs and build folders
    ignore_dirs = {".git", "node_modules", "venv", ".venv", "__pycache__", "build", "dist", ".next"}
    
    all_files = []
    for root, dirs, files in os.walk(folder_path):
        dirs[:] = [d for d in dirs if d not in ignore_dirs]
        for f in files:
            all_files.append(os.path.join(root, f))

    total_files = len(all_files)
    print(f"[Report Generator] Found {total_files} total files to categorize.", flush=True)

    # 1. Statistics & Extensions
    stats = {
        "Python (.py)": 0,
        "JavaScript/TypeScript (.js, .jsx, .ts, .tsx)": 0,
        "Stylesheets (.css, .scss)": 0,
        "Web Pages (.html)": 0,
        "Documents (.pdf, .txt, .md, .docx, .csv)": 0,
        "Images (.png, .jpg, .jpeg, .gif, .webp, .bmp)": 0,
        "Log Files (.log)": 0,
        "Others": 0
    }

    # Lists for specific processors
    python_files = []
    js_ts_files = []
    image_files = []
    log_files = []
    document_files = []

    for file_path in all_files:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".py":
            stats["Python (.py)"] += 1
            python_files.append(file_path)
        elif ext in (".js", ".jsx", ".ts", ".tsx"):
            stats["JavaScript/TypeScript (.js, .jsx, .ts, .tsx)"] += 1
            js_ts_files.append(file_path)
        elif ext in (".css", ".scss"):
            stats["Stylesheets (.css, .scss)"] += 1
        elif ext in (".html", ".htm"):
            stats["Web Pages (.html)"] += 1
        elif ext in (".txt", ".pdf", ".md", ".docx", ".csv"):
            stats["Documents (.pdf, .txt, .md, .docx, .csv)"] += 1
            document_files.append(file_path)
        elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"):
            stats["Images (.png, .jpg, .jpeg, .gif, .webp, .bmp)"] += 1
            image_files.append(file_path)
        elif ext == ".log" or "log" in os.path.basename(file_path).lower():
            stats["Log Files (.log)"] += 1
            log_files.append(file_path)
        else:
            stats["Others"] += 1

    # 2. Code Analysis (AST & Framework Detection)
    detected_frameworks = set()
    classes = []
    functions = []

    # Parse first 5 python files for classes and functions
    if python_files:
        print("[FileMind AI] 📂 Extracting Python class/function structures via AST parsing...", flush=True)
    for py_file in python_files[:5]:
        try:
            with open(py_file, "r", encoding="utf-8") as f:
                code = f.read()
            tree = ast.parse(code)
            classes.extend([node.name for node in ast.walk(tree) if isinstance(node, ast.ClassDef)])
            functions.extend([node.name for node in ast.walk(tree) if isinstance(node, ast.FunctionDef)])
            
            # Simple Python framework detection
            if "flask" in code.lower():
                detected_frameworks.add("Flask")
            if "django" in code.lower():
                detected_frameworks.add("Django")
            if "fastapi" in code.lower():
                detected_frameworks.add("FastAPI")
        except:
            pass

    # Scan JS/TS files for frameworks
    if js_ts_files:
        print("[FileMind AI] 🌐 Scanning JS/TS modules for active frameworks...", flush=True)
    for js_file in js_ts_files[:5]:
        try:
            with open(js_file, "r", encoding="utf-8") as f:
                code = f.read()
            if "react" in code.lower():
                detected_frameworks.add("React")
            if "express" in code.lower():
                detected_frameworks.add("Express")
            if "next" in code.lower() or "next/router" in code.lower():
                detected_frameworks.add("Next.js")
            if "vue" in code.lower():
                detected_frameworks.add("Vue.js")
        except:
            pass

    # 3. OCR Results
    ocr_results = {}
    if PADDLE_AVAILABLE and image_files:
        print("[FileMind AI] 👁️ Performing OCR on discovered image assets...", flush=True)
        try:
            ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
            for img in image_files[:3]:
                result = ocr.ocr(img, cls=True)
                extracted = []
                if result and result[0]:
                    for line in result[0]:
                        extracted.append(line[1][0])
                ocr_results[os.path.basename(img)] = "\n".join(extracted) if extracted else "[No text found]"
        except Exception as e:
            print(f"[Warning] OCR failed: {e}", flush=True)
    elif image_files:
        for img in image_files[:3]:
            ocr_results[os.path.basename(img)] = "[PaddleOCR not loaded or installed. OCR skipped.]"

    # 4. Document Summaries
    doc_summaries = {}
    if document_files:
        print("[FileMind AI] 📝 Reading and summarizing documentation assets...", flush=True)
    for doc in document_files[:3]:
        # Read text documents or markdown to summarize
        try:
            content = read_file(doc)
            lines = content.splitlines()[:50]  # Read first 50 lines
            snippet = "\n".join(lines)
            summary_prompt = f"Please summarize the purpose of this document snippet in 2-3 sentences:\n\nDocument: {os.path.basename(doc)}\nSnippet:\n{snippet}"
            summary = ask_ai(summary_prompt, save_to_memory=False, include_history=False)
            doc_summaries[os.path.basename(doc)] = summary
        except Exception as e:
            doc_summaries[os.path.basename(doc)] = f"Could not summarize: {e}"

    # 5. Log Analysis
    log_warnings_errors = {}
    if log_files:
        print("[FileMind AI] ⚙️ Auditing system and server log registers...", flush=True)
    for log in log_files[:3]:
        try:
            with open(log, "r", encoding="utf-8") as f:
                lines = f.readlines()
            issues = []
            for line in lines:
                if any(kw in line.upper() for kw in ("ERROR", "WARN", "FAIL", "CRITICAL")):
                    issues.append(line.strip())
            log_warnings_errors[os.path.basename(log)] = issues[:10]  # Limit to first 10 matches
        except Exception as e:
            log_warnings_errors[os.path.basename(log)] = [f"Failed to read log file: {e}"]

    # 6. AI Final Master Summary
    print("[FileMind AI] 🧠 Requesting executive AI Master Project Summary...", flush=True)
    ai_summary_prompt = f"""You are the master report generator for FileMind. Please provide a professional, executive AI Project Summary of the project based on the following scanned structural details:

Project Name: {project_name}
Total Files: {total_files}
File Categories:
{stats}

Detected Frameworks:
{list(detected_frameworks)}

Detected Classes (sample):
{classes[:15]}

Detected Functions (sample):
{functions[:15]}

Provide a clean, elegant, high-level overview explaining what this project likely is, its architecture, and recommended next steps or improvements.
"""
    final_ai_summary = ask_ai(ai_summary_prompt, save_to_memory=False, include_history=False)

    # 7. Compile Clean Markdown Report
    analysis_date = datetime.now().strftime("%d-%m-%Y")
    timestamp_filename = datetime.now().strftime("%Y-%m-%d")
    
    report_md = f"""# FileMind AI - Deep Project Analysis Report

## 1. Project Information
* **Project Name:** `{project_name}`
* **Analysis Date:** {analysis_date}
* **Total Scanned Files:** {total_files}

---

## 2. File Statistics
| File Type | Count |
| :--- | :--- |
"""
    for cat, count in stats.items():
        report_md += f"| {cat} | {count} |\n"

    report_md += """
---

## 3. Code Analysis
"""
    if detected_frameworks:
        report_md += "* **Detected Frameworks:**\n"
        for fw in detected_frameworks:
            report_md += f"  - {fw}\n"
    else:
        report_md += "* **Detected Frameworks:** None identified.\n"

    if functions:
        report_md += "\n* **Key Functions (Scanned):**\n"
        for func in list(set(functions))[:10]:
            report_md += f"  - `{func}()`\n"
            
    if classes:
        report_md += "\n* **Key Classes (Scanned):**\n"
        for cls in list(set(classes))[:10]:
            report_md += f"  - `{cls}`\n"

    report_md += """
---

## 4. OCR Results
"""
    if ocr_results:
        for filename, text in ocr_results.items():
            report_md += f"### File: `{filename}`\n"
            report_md += f"```text\n{text}\n```\n\n"
    else:
        report_md += "*No image files were scanned.*\n"

    report_md += """
---

## 5. Document Summaries
"""
    if doc_summaries:
        for filename, summary in doc_summaries.items():
            report_md += f"### File: `{filename}`\n"
            report_md += f"{summary}\n\n"
    else:
        report_md += "*No documents were scanned.*\n"

    report_md += """
---

## 6. Log Analysis
"""
    if log_warnings_errors:
        for filename, issues in log_warnings_errors.items():
            report_md += f"### File: `{filename}`\n"
            if issues:
                report_md += "* **Issues Found:**\n"
                for issue in issues:
                    report_md += f"  - `{issue}`\n"
            else:
                report_md += "* No immediate Errors or Warnings detected in log file.*\n"
    else:
        report_md += "*No log files were scanned.*\n"

    report_md += f"""
---

## 7. AI Final Summary
{final_ai_summary}

---
*Report compiled natively on local CPU by FileMind AI.*
"""

    # 8. Save the Report
    os.makedirs(save_directory, exist_ok=True)
    report_filename = os.path.join(save_directory, f"analysis_report_{timestamp_filename}.{fmt}")
    project_report_filename = os.path.join(save_directory, f"project_report.{fmt}")

    if fmt == "docx":
        try:
            import docx
            doc = docx.Document()
            doc.add_heading("FileMind AI - Deep Project Analysis Report", 0)
            
            lines = report_md.split('\n')
            for line in lines:
                if "FileMind AI - Deep Project Analysis Report" in line: continue
                clean_line = line.replace('**', '').replace('`', '')
                if clean_line.startswith('### '):
                    doc.add_heading(clean_line[4:], 3)
                elif clean_line.startswith('## '):
                    doc.add_heading(clean_line[3:], 2)
                elif clean_line.startswith('# '):
                    doc.add_heading(clean_line[2:], 1)
                elif clean_line.startswith('- ') or clean_line.startswith('* '):
                    doc.add_paragraph(clean_line[2:], style='List Bullet')
                elif clean_line.strip() != "":
                    doc.add_paragraph(clean_line)
                    
            doc.save(report_filename)
            doc.save(project_report_filename)
            
            print(f"\n[Report Generator] Success! Dynamic report written to:")
            print(f"  - {report_filename}")
            print(f"  - {project_report_filename}")
            return report_filename
        except ImportError:
            print("\n[Warning] 'python-docx' package not installed. Could not generate Word Document. Please run 'pip install python-docx'.")
            return None
    else:
        if fmt == "html":
            try:
                import markdown
                html_body = markdown.markdown(report_md)
                output_content = f"<html><head><style>body{{font-family:sans-serif;line-height:1.6;padding:40px;max-width:900px;margin:auto;table{{border-collapse:collapse;width:100%;}}th,td{{border:1px solid #ddd;padding:8px;}}th{{background-color:#f2f2f2;}}</style></head><body>{html_body}</body></html>"
            except ImportError:
                output_content = report_md
                print("\n[Warning] 'markdown' package not installed. HTML report contains raw markdown.")
        elif fmt == "txt":
            output_content = report_md.replace("#", "").replace("**", "").replace("`", "")
        else:
            output_content = report_md
            
        # Write specific dated report
        with open(report_filename, "w", encoding="utf-8") as f:
            f.write(output_content)
    
        # Overwrite/Create project_report as the active master report
        with open(project_report_filename, "w", encoding="utf-8") as f:
            f.write(output_content)
    
        print(f"\n[Report Generator] Success! Dynamic report written to:")
        print(f"  - {report_filename}")
        print(f"  - {project_report_filename}")
    
        return report_filename
