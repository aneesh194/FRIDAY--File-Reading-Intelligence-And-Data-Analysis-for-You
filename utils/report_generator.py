import os
from datetime import datetime


def save_report(analysis_type, target_path, content, save_directory="reports", fmt="md"):
    """
    Saves a beautifully structured analysis report to the specified directory in the requested format,
    and appends a summary of the analysis to reports/project_report.txt.
    """
    try:
        os.makedirs(save_directory, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        target_name = os.path.basename(target_path.replace("\\", "/")).replace(" ", "_")
        if not target_name:
            target_name = "analysis"
            
        report_filename = os.path.join(save_directory, f"analysis_{analysis_type}_{target_name}_{timestamp}.{fmt}")
        readable_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Format the markdown content
        report_md = f"""# FileMind AI - Analysis Report
**Generated on:** {readable_time}
**Target Analyzed:** `{target_path}`
**Analysis Type:** {analysis_type.upper()}

---

## Analysis Details

{content}

---
*Report generated automatically by FileMind AI.*
"""

        if fmt == "docx":
            try:
                import docx
                doc = docx.Document()
                doc.add_heading("FileMind AI - Analysis Report", 0)
                doc.add_paragraph(f"Generated on: {readable_time}")
                doc.add_paragraph(f"Target Analyzed: {target_path}")
                doc.add_paragraph(f"Analysis Type: {analysis_type.upper()}")
                doc.add_heading("Analysis Details", 1)
                
                # Parse markdown-like content to word paragraphs
                lines = content.split('\n')
                for line in lines:
                    clean_line = line.replace('**', '').replace('`', '')
                    if clean_line.startswith('### '):
                        doc.add_heading(clean_line[4:], 3)
                    elif clean_line.startswith('## '):
                        doc.add_heading(clean_line[3:], 2)
                    elif clean_line.startswith('# '):
                        doc.add_heading(clean_line[2:], 1)
                    elif clean_line.startswith('- ') or clean_line.startswith('* '):
                        doc.add_paragraph(clean_line[2:], style='List Bullet')
                    elif clean_line.strip() != "":
                        doc.add_paragraph(clean_line)
                        
                doc.add_paragraph("\nReport generated automatically by FileMind AI.")
                doc.save(report_filename)
            except ImportError:
                print("\n[Warning] 'python-docx' package not installed. Could not generate Word Document. Please run 'pip install python-docx'.")
                return
        else:
            if fmt == "html":
                try:
                    import markdown
                    html_body = markdown.markdown(report_md)
                    output_content = f"<html><head><style>body{{font-family:sans-serif;line-height:1.6;padding:40px;max-width:900px;margin:auto;}}</style></head><body>{html_body}</body></html>"
                except ImportError:
                    output_content = report_md
                    print("\n[Warning] 'markdown' package not installed. HTML report contains raw markdown.")
            elif fmt == "txt":
                output_content = report_md.replace("#", "").replace("**", "").replace("`", "")
            else:
                output_content = report_md
                
            # Write the report file
            with open(report_filename, "w", encoding="utf-8") as f:
                f.write(output_content)
            
        print(f"\n[Report saved successfully to: {report_filename}]")
        
        # Append to the master reports/project_report.txt
        master_report_path = "reports/project_report.txt"
        os.makedirs("reports", exist_ok=True)
        
        # Create master file header if it doesn't exist
        if not os.path.exists(master_report_path):
            with open(master_report_path, "w", encoding="utf-8") as master_f:
                master_f.write("=========================================\n")
                master_f.write("        FILEMIND AI PROJECT REPORT       \n")
                master_f.write("=========================================\n\n")
        
        with open(master_report_path, "a", encoding="utf-8") as master_f:
            master_f.write(f"[{readable_time}] TYPE: {analysis_type.upper()} | TARGET: {target_path}\n")
            master_f.write(f"Saved Report: {report_filename}\n")
            master_f.write("-" * 50 + "\n\n")
            
    except Exception as e:
        print(f"\n[Error saving report: {e}]")
